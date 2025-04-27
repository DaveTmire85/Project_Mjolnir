from utils.text_cleaner import clean_text
from db.db_manager import insert_entry
from docx import Document

class SpellsParser:
    def __init__(self):
        self.class_mapping = {}
        self.sorted_mapping = {}
        self.description_mapping = {}

    def parse(self, file_paths):
        spell_list_path, sorted_path, descriptions_path = file_paths
        
        self.parse_spell_list(spell_list_path)
        self.parse_spells_sorted(sorted_path)
        self.parse_spell_descriptions(descriptions_path)

        # Merge all data and insert
        for spell_name in self.description_mapping:
            spell = self.description_mapping[spell_name]
            spell['classes'] = self.class_mapping.get(spell_name, {})
            spell['domains'] = self.sorted_mapping.get(spell_name, {}).get('domains', [])
            spell['sourcebooks'] = self.sorted_mapping.get(spell_name, {}).get('sources', [])
            
            insert_entry('spells', spell)

    def parse_spell_list(self, file_path):
        doc = Document(file_path)
        lines = clean_text([p.text for p in doc.paragraphs])

        current_class = None
        current_level = None

        for line in lines:
            if "Spell List" in line:
                current_class = line.split("Spell List")[0].strip()
            elif "Level" in line:
                parts = line.split("Level")
                if len(parts) >= 2:
                    current_level = parts[0].strip()
            elif line:
                spell_name = line.split("(")[0].strip()
                if spell_name:
                    if spell_name not in self.class_mapping:
                        self.class_mapping[spell_name] = {}
                    self.class_mapping[spell_name][current_class] = current_level

    def parse_spells_sorted(self, file_path):
        doc = Document(file_path)
        lines = clean_text([p.text for p in doc.paragraphs])

        for line in lines:
            if "]" in line and "(" not in line[:10]:
                # Format: SpellName [ClassLevelInfo]
                name_part = line.split("[")[0].strip()
                details_part = line.split("[")[1].split("]")[0].strip()
                domains = []
                sources = []
                
                if name_part not in self.sorted_mapping:
                    self.sorted_mapping[name_part] = {}
                
                for tag in details_part.split():
                    if tag.isupper() and len(tag) > 2:
                        domains.append(tag)
                    else:
                        sources.append(tag)
                
                self.sorted_mapping[name_part]['domains'] = domains
                self.sorted_mapping[name_part]['sources'] = sources

    def parse_spell_descriptions(self, file_path):
        doc = Document(file_path)
        lines = clean_text([p.text for p in doc.paragraphs])

        current_spell = None
        for line in lines:
            if "<" in line and ">" in line:
                # New spell start
                if current_spell:
                    self.description_mapping[current_spell['name']] = current_spell
                parts = line.split("<")
                name = parts[0].strip()
                mechanics = parts[1].split(">")[0].strip()
                current_spell = {
                    'name': name,
                    'mechanics': mechanics,
                    'description': ""
                }
            elif current_spell:
                current_spell['description'] += line + " "

        if current_spell:
            self.description_mapping[current_spell['name']] = current_spell
