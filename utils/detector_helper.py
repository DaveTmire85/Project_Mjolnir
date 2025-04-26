import docx

def load_text_for_detection(file_path):
    """
    Extracts all text content from a .docx file.
    
    Args:
        file_path (str): Path to the .docx file
        
    Returns:
        list: All text content from paragraphs and tables
        
    Raises:
        FileNotFoundError: If the file doesn't exist
    """
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return text
